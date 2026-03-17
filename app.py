from flask import Flask, render_template, request, send_file
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from dateutil.relativedelta import relativedelta
from num2words import num2words

app = Flask(__name__)


# -----------------------------
# Indian currency formatting
# -----------------------------
def format_inr(number):

    number = int(round(number))
    s = str(number)

    if len(s) <= 3:
        return s

    last3 = s[-3:]
    rest = s[:-3]

    parts = []

    while len(rest) > 2:
        parts.insert(0, rest[-2:])
        rest = rest[:-2]

    if rest:
        parts.insert(0, rest)

    return ",".join(parts + [last3])


# -----------------------------
# Font formatting
# -----------------------------
def format_cell(cell, align="left", bold=False):

    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()

    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold

    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# -----------------------------
# Table borders
# -----------------------------
def set_cell_border(cell):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    borders = OxmlElement('w:tcBorders')

    for border in ('top','left','bottom','right'):
        edge = OxmlElement(f'w:{border}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '8')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), '000000')
        borders.append(edge)

    tcPr.append(borders)


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():

    start_date = request.form["start_date"]
    lease_years = int(request.form["lease_years"])
    total_sft = float(request.form["total_sft"])
    rate_per_sft = float(request.form["rate_per_sft"])
    enhancement = float(request.form["rent_enhancement"])
    advance_amount = int(request.form["advance_amount"])
    whitewash_years = request.form["white_wash_years"]

    start = datetime.strptime(start_date,"%Y-%m-%d")
    end = start + relativedelta(years=lease_years)

    base_rent = int(total_sft * rate_per_sft)

    monthly = base_rent
    blocks = lease_years // 2

    rows = []
    total_sum = 0

    # -----------------------------
    # Rent progression table
    # -----------------------------
    for i in range(blocks):

        block_start = start + relativedelta(years=2*i)
        block_end = block_start + relativedelta(years=2) - relativedelta(days=1)

        total_24 = monthly * 24
        total_sum += total_24

        if i == 0:
            label = "1st Two Years"
        elif i == blocks-1:
            label = "Last Two Years"
        else:
            label = "Next Two Years"

        rows.append({
            "sno": i+1,
            "period": f"{block_start.strftime('%d-%m-%Y')} to {block_end.strftime('%d-%m-%Y')}",
            "label": label,
            "monthly": format_inr(monthly),
            "total": format_inr(total_24)
        })

        monthly = int(monthly + (monthly * enhancement/100))

    avg_monthly_rent = total_sum / (lease_years * 12)
    avg_rent = total_sum / (lease_years)

    round_rent = round(avg_rent)

    total_cost_words = num2words(base_rent, lang="en_IN").title()
    advance_words = num2words(advance_amount, lang="en_IN").title()

    # -----------------------------
    # Fill text placeholders
    # -----------------------------
    doc = DocxTemplate("lease_template.docx")

    context = {

        "START_DATE": start.strftime("%d-%m-%Y"),
        "END_DATE": end.strftime("%d-%m-%Y"),

        "LEASE_YEARS": lease_years,

        "TOTAL_SFT": total_sft,
        "RATE_PER_SFT": rate_per_sft,

        "TOTAL_COST": format_inr(base_rent),
        "TOTAL_COST_WORDS": total_cost_words,

        "RENT_ENHANCEMENT": enhancement,

        "ADVANCE_AMOUNT": format_inr(advance_amount),
        "ADVANCE_AMOUNT_WORDS": advance_words,

        "WHITE_WASH_YEARS": whitewash_years,

        "AVG_RENT": format_inr(avg_rent),
        "AVG_MONTHLY_RENT": format_inr(avg_monthly_rent),

        "ROUND_RENT": format_inr(round_rent),

        "TOTAL_MONTHS": lease_years * 12,
        "TOTAL_SUM": format_inr(total_sum)
    }

    doc.render(context)

    temp_file = "temp.docx"
    doc.save(temp_file)

    # -----------------------------
    # Insert rows into table
    # -----------------------------
    document = Document(temp_file)

    for table in document.tables:

        for i, row in enumerate(table.rows):

            if row.cells[0].text.strip() == "S.No":

                header_index = i

                for r in rows:

                    new_row = table.add_row()
                    cells = new_row.cells

                    cells[0].text = str(r["sno"])
                    cells[1].text = r["period"]
                    cells[2].text = r["label"]
                    cells[3].text = r["monthly"]
                    cells[4].text = r["total"]

                    format_cell(cells[0],"center")
                    format_cell(cells[1])
                    format_cell(cells[2])
                    format_cell(cells[3],"center")
                    format_cell(cells[4],"center")

                    for c in cells:
                        set_cell_border(c)

                # -----------------------------
                # Summary row
                # -----------------------------
                total_row = table.add_row()
                cells = total_row.cells

                merged = cells[0].merge(cells[1]).merge(cells[2])
                merged.text = f"Total For {lease_years*12} Months"

                cells[3].text = format_inr(avg_monthly_rent)
                cells[4].text = format_inr(total_sum)

                format_cell(merged,"center",True)
                format_cell(cells[3],"center",True)
                format_cell(cells[4],"center",True)

                for c in cells:
                    set_cell_border(c)

                break

    output_file = "lease_generated.docx"

    document.save(output_file)

    return send_file(output_file, as_attachment=True)


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
